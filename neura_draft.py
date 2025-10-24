#=========================================================================================================#
# INSTALL THE REQUIRED PACKAGES , USE COMMAND:                                                            #
# pip install streamlit langchain langchain-google-genai fpdf google-generativeai python-docx markdown re #
#=========================================================================================================#
import streamlit as st
import os
import re
import tempfile
from datetime import datetime
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.schema import HumanMessage, SystemMessage
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
import markdown


class MultiFormatDocumentGenerator:
    def __init__(self):
        self.pdf = FPDF()
        self.pdf.set_auto_page_break(auto=True, margin=15)
        self.pdf.set_margins(left=20, top=20, right=20)

    def clean_text(self, text):
        """Remove special characters and clean text for documents"""
        if not text:
            return ""

        # Replace common Unicode characters with ASCII equivalents
        replacements = {
            "\u2013": "-",  # en dash
            "\u2014": "-",  # em dash
            "\u2018": "'",  # left single quotation mark
            "\u2019": "'",  # right single quotation mark
            "\u201c": '"',  # left double quotation mark
            "\u201d": '"',  # right double quotation mark
            "\u2022": "*",  # bullet
            "\u2026": "...",  # ellipsis
            "\u00a0": " ",  # non-breaking space
            "\u00b0": " degrees ",  # degree symbol
            "\u00ae": "(R)",  # registered trademark
            "\u00a9": "(C)",  # copyright
            "\u2122": "(TM)",  # trademark
        }

        for unicode_char, ascii_char in replacements.items():
            text = text.replace(unicode_char, ascii_char)

        text = re.sub(r"#+\s*", "", text)
        text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
        text = re.sub(r"\*(.*?)\*", r"\1", text)
        text = re.sub(r"`(.*?)`", r"\1", text)
        text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)

        text = re.sub(r"\n\s*\n", "\n\n", text)
        text = text.strip()

        return text

    def safe_encode(self, text):
        """Safely encode text for PDF generation"""
        if not text:
            return ""

        return text.encode("latin-1", "replace").decode("latin-1")

    def create_pdf(self, title, outline, content, output_path):
        """Create a fully formatted PDF document with Unicode support"""
        try:
            self.pdf = FPDF()
            self.pdf.set_auto_page_break(auto=True, margin=15)
            self.pdf.set_margins(left=20, top=20, right=20)

            self.pdf.add_page()

            clean_title = self.safe_encode(self.clean_text(title))

            # Title only - no date
            self.pdf.set_font("Arial", style="B", size=18)
            self.pdf.set_text_color(0, 51, 102)
            self.pdf.cell(0, 15, txt=clean_title, ln=1, align="C")
            self.pdf.ln(10)  

            sections = self.parse_content_into_sections(content)
            for section_title, section_content in sections:
                clean_section_title = self.safe_encode(self.clean_text(section_title))
                self.add_section(
                    clean_section_title,
                    section_content,
                    level=1 if "#" in section_title else 2,
                )
                self.pdf.ln(5)

            self.pdf.output(output_path)
            return output_path

        except Exception as e:
            st.error(f"Error creating PDF: {str(e)}")
            return self.create_simple_pdf(title, content, output_path)

    def create_simple_pdf(self, title, content, output_path):
        """Fallback PDF creation with minimal formatting"""
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            pdf.cell(200, 10, txt=self.safe_encode(title), ln=1, align="C")
            pdf.ln(10)

            clean_content = self.safe_encode(self.clean_text(content))
            pdf.multi_cell(0, 10, txt=clean_content)

            pdf.output(output_path)
            return output_path
        except Exception as e:
            st.error(f"Error in fallback PDF creation: {str(e)}")
            return None

    def add_section(self, title, content, level=1):
        """Add a section with proper formatting"""
        try:
            if level == 1:
                self.pdf.set_font("Arial", "B", 14)
                self.pdf.set_text_color(0, 51, 102)
                self.pdf.cell(0, 10, txt=title, ln=1)
                self.pdf.ln(5)
            elif level == 2:
                self.pdf.set_font("Arial", "B", 12)
                self.pdf.set_text_color(0, 0, 0)
                self.pdf.cell(0, 8, txt=title, ln=1)
                self.pdf.ln(3)

            self.pdf.set_font("Arial", size=11)
            self.pdf.set_text_color(0, 0, 0)

            cleaned_content = self.safe_encode(self.clean_text(content))
            paragraphs = cleaned_content.split("\n\n")

            for paragraph in paragraphs:
                if paragraph.strip():

                    if paragraph.strip().startswith(("-", "•", "*")):
                        lines = paragraph.split("\n")
                        for line in lines:
                            if line.strip():
                                self.pdf.set_font("Arial", size=11)
                                self.pdf.cell(10)
                                clean_line = self.safe_encode(line.lstrip("-•* "))
                                self.pdf.multi_cell(0, 6, txt=f"• {clean_line}")
                    else:
                        self.pdf.multi_cell(0, 6, txt=paragraph)
                    self.pdf.ln(3)
        except Exception as e:
            st.warning(f"Could not add section '{title}': {str(e)}")

    def create_docx(self, title, outline, content, output_path):
        """Create a Microsoft Word document with Unicode support"""
        try:
            doc = Document()

            title_para = doc.add_heading(self.clean_text(title), 0)
            title_para.alignment = 1

            sections = self.parse_content_into_sections(content)
            for section_title, section_content in sections:

                clean_section_title = self.clean_text(section_title)
                doc.add_heading(
                    clean_section_title, level=1 if len(clean_section_title) < 50 else 2
                )

                cleaned_content = self.clean_text(section_content)
                paragraphs = cleaned_content.split("\n\n")

                for paragraph in paragraphs:
                    if paragraph.strip():
                        if paragraph.strip().startswith(("-", "•", "*")):

                            lines = paragraph.split("\n")
                            for line in lines:
                                if line.strip():
                                    clean_line = self.clean_text(line.lstrip("-•* "))
                                    doc.add_paragraph(
                                        f"• {clean_line}", style="List Bullet"
                                    )
                        else:
                            doc.add_paragraph(paragraph)

                doc.add_paragraph()

            doc.save(output_path)
            return output_path
        except Exception as e:
            st.error(f"Error creating DOCX: {str(e)}")
            return None

    def create_txt(self, title, outline, content, output_path):
        """Create a plain text document with Unicode support"""
        try:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(f"{self.clean_text(title)}\n")
                f.write("=" * len(title) + "\n\n")
                
                cleaned_content = self.clean_text(content)
                f.write(cleaned_content)

            return output_path
        except Exception as e:
            st.error(f"Error creating TXT: {str(e)}")
            return None

    def create_html(self, title, outline, content, output_path):
        """Create an HTML document with Unicode support"""
        try:
            html_content = f"""
            <!DOCTYPE html>
            <html lang="en">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>{self.clean_text(title)}</title>
                <style>
                    body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; color: #333; }}
                    h1 {{ color: #003366; text-align: center; border-bottom: 2px solid #003366; padding-bottom: 10px; }}
                    h2 {{ color: #003366; border-left: 4px solid #003366; padding-left: 10px; }}
                    h3 {{ color: #555; }}
                    .meta {{ text-align: center; color: #666; font-style: italic; margin-bottom: 30px; }}
                    .toc {{ background: #f9f9f9; padding: 20px; border-radius: 5px; margin: 20px 0; }}
                    .footer {{ margin-top: 40px; padding-top: 20px; border-top: 1px solid #ddd; text-align: center; color: #666; font-size: 0.9em; }}
                    ul {{ margin: 10px 0; }}
                    li {{ margin: 5px 0; }}
                </style>
            </head>
            <body>
                <h1>{self.clean_text(title)}</h1>
            """

            html_content += '<div class="content">\n'

            try:
                html_content += markdown.markdown(content)
            except:
                html_content += f"<pre>{self.clean_text(content)}</pre>"

            html_content += "\n</div>"
            html_content += "\n</body>\n</html>"

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(html_content)

            return output_path
        except Exception as e:
            st.error(f"Error creating HTML: {str(e)}")
            return None

    def parse_content_into_sections(self, content):
        """Parse content into sections based on headings"""
        sections = []
        if not content:
            return sections

        lines = content.split("\n")
        current_section = []
        current_title = "Introduction"

        for line in lines:
            if line.strip().startswith("## "):
                if current_section:
                    sections.append((current_title, "\n".join(current_section)))
                    current_section = []
                current_title = line.replace("##", "").strip()
            elif line.strip().startswith("### "):
                if current_section:
                    sections.append((current_title, "\n".join(current_section)))
                    current_section = []
                current_title = line.replace("###", "").strip()
            elif line.strip():
                current_section.append(line)

        if current_section:
            sections.append((current_title, "\n".join(current_section)))

        return sections


class AgenticDocumentGenerator:
    def __init__(self):
        self.llm = None
        self.conversation_history = []

    def initialize_llm(self, api_key, model_name="gemini-pro"):
        """Initialize the Gemini LLM with the provided API key and model"""
        try:
            self.llm = ChatGoogleGenerativeAI(
                model=model_name, google_api_key=api_key, temperature=0.7
            )
            return True
        except Exception as e:
            st.error(f"Error initializing Gemini with model {model_name}: {str(e)}")
            return False

    def generate_outline(self, topic, doc_type="Article"):
        """Generate a detailed outline for the document"""
        try:
            if doc_type == "Report":
                prompt = f"""
                Create a comprehensive outline for a detailed report on: {topic}
                
                The outline should include:
                1. Executive Summary
                2. Introduction and Background
                3. Main Analysis Sections (3-5 sections with subsections)
                4. Case Studies/Examples
                5. Findings and Recommendations
                6. Conclusion
                7. References
                
                Please provide a structured outline with clear sections and brief descriptions of what each section will cover.
                """
            elif doc_type == "Research Paper":
                prompt = f"""
                Create a comprehensive outline for a research paper on: {topic}
                
                The outline should include:
                1. Abstract
                2. Introduction
                3. Literature Review
                4. Methodology
                5. Results
                6. Discussion
                7. Conclusion
                8. References
                
                Please provide a structured outline with clear sections and brief descriptions.
                """
            else:
                prompt = f"""
                Create a comprehensive outline for an article on: {topic}
                
                The outline should include:
                1. Introduction
                2. Historical Context/Background
                3. Key Concepts and Definitions
                4. Current State/Applications
                5. Challenges and Opportunities
                6. Future Outlook
                7. Conclusion
                
                Please provide a structured outline with clear sections and brief descriptions of what each section will cover.
                """

            messages = [
                SystemMessage(
                    content="You are an expert technical writer. Create detailed, logical outlines for documents."
                ),
                HumanMessage(content=prompt),
            ]

            response = self.llm(messages)
            self.conversation_history.append(("outline", response.content))
            return response.content

        except Exception as e:
            st.error(f"Error generating outline: {str(e)}")
            return None

    def generate_content_based_on_outline(self, topic, outline, user_feedback=""):
        """Generate full content based on the approved outline"""
        try:
            prompt = f"""
            Based on the following topic and approved outline, generate a comprehensive, well-formatted document.
            
            TOPIC: {topic}
            
            APPROVED OUTLINE:
            {outline}
            
            USER FEEDBACK/REQUIREMENTS:
            {user_feedback}
            
            Requirements:
            1. Follow the outline structure exactly
            2. Write detailed, informative content for each section
            3. Use proper markdown formatting with ## for main headings and ### for subheadings
            4. Include relevant examples, data, and explanations
            5. Ensure the content flows logically between sections
            6. Make it engaging and educational
            7. Each section should be substantial (3-5 paragraphs minimum for main sections)
            
            Please generate the complete document now.
            """

            messages = [
                SystemMessage(
                    content="You are a professional content writer. Create comprehensive, well-structured documents with proper formatting."
                ),
                HumanMessage(content=prompt),
            ]

            response = self.llm(messages)
            self.conversation_history.append(("content", response.content))
            return response.content

        except Exception as e:
            st.error(f"Error generating content: {str(e)}")
            return None


def main():
    st.set_page_config(
        page_title="AI Document Generator - Multi Format", page_icon="📄", layout="wide"
    )

    st.title("AI Document Generator - Multi-Format Download")
    st.markdown(
        "Generate professional documents in multiple formats through an interactive, multi-step process"
    )

    if "current_step" not in st.session_state:
        st.session_state.current_step = 1
    if "outline" not in st.session_state:
        st.session_state.outline = None
    if "final_content" not in st.session_state:
        st.session_state.final_content = None
    if "topic" not in st.session_state:
        st.session_state.topic = ""
    if "doc_type" not in st.session_state:
        st.session_state.doc_type = "Article"
    if "user_feedback" not in st.session_state:
        st.session_state.user_feedback = ""
    if "selected_model" not in st.session_state:
        st.session_state.selected_model = "gemini-pro"

    with st.sidebar:
        st.header("Configuration")

        api_key = st.text_input(
            "Enter your Gemini API Key:",
            type="password",
            help="Get your API key from: https://aistudio.google.com/app/apikey",
        )

        # Model selection
        st.subheader("Model Selection")
        model_options = {
            "Gemini Pro": "gemini-pro",
            "Gemini Pro Vision": "gemini-pro-vision",
            "Gemini 2.5 Flash": "gemini-2.5-flash",
        }

        selected_model = st.selectbox(
            "Choose Gemini Model:",
            options=list(model_options.keys()),
            index=0,
            help="Gemini Pro for text, Gemini Pro Vision for image understanding",
        )
        st.session_state.selected_model = model_options[selected_model]

        if api_key:
            st.session_state.api_key = api_key
            st.success("API Key saved!")

        st.markdown("---")
        st.header("Progress")

        steps = [
            "1. Enter Topic & Type",
            "2. Review Outline",
            "3. Provide Feedback",
            "4. Generate Content",
            "5. Download Documents",
        ]

        for i, step in enumerate(steps, 1):
            if i == st.session_state.current_step:
                st.markdown(f"**{step}**")
            elif i < st.session_state.current_step:
                st.markdown(f"{step}")
            else:
                st.markdown(f"{step}")

        st.markdown("---")

        if st.button("Reset Workflow"):
            for key in list(st.session_state.keys()):
                if key != "api_key":
                    del st.session_state[key]
            st.session_state.current_step = 1
            st.rerun()

    if not st.session_state.get("api_key"):
        st.warning("Please enter your Gemini API key in the sidebar to continue.")
        return

    doc_gen = AgenticDocumentGenerator()
    if not doc_gen.initialize_llm(
        st.session_state.api_key, st.session_state.selected_model
    ):
        return

    if st.session_state.current_step == 1:
        st.header("Step 1: Define Your Document")

        col1, col2 = st.columns([2, 1])

        with col1:
            topic = st.text_input(
                "Enter your topic:",
                value=st.session_state.topic,
                placeholder="e.g., Artificial Intelligence in Healthcare, Renewable Energy Solutions...",
                help="Be specific about what you want to learn or document",
            )

        with col2:
            doc_type = st.selectbox(
                "Document Type:",
                ["Article", "Report", "Research Paper", "Guide"],
                index=0 if st.session_state.doc_type == "Article" else 1,
            )

        if st.button("Generate Outline →", type="primary") and topic:
            st.session_state.topic = topic
            st.session_state.doc_type = doc_type

            with st.spinner("Generating detailed outline..."):
                outline = doc_gen.generate_outline(topic, doc_type)
                if outline:
                    st.session_state.outline = outline
                    st.session_state.current_step = 2
                    st.rerun()
    elif st.session_state.current_step == 2:
        st.header("Step 2: Review & Approve Outline")

        st.subheader(f"Outline for: {st.session_state.topic}")
        st.info(
            "Review the outline below. You can provide feedback or request modifications."
        )

        with st.expander("Generated Outline", expanded=True):
            st.markdown(st.session_state.outline)

        st.subheader("Provide Feedback (Optional)")
        feedback = st.text_area(
            "Any modifications or specific requirements?",
            placeholder="e.g., Add more sections about X, Focus more on Y, Include examples about Z...",
            height=100,
            key="feedback_input",
        )

        col1, col2 = st.columns(2)

        with col1:
            if st.button("← Back to Topic", type="secondary"):
                st.session_state.current_step = 1
                st.rerun()

        with col2:
            if st.button("Generate Final Document →", type="primary"):
                st.session_state.user_feedback = feedback
                st.session_state.current_step = 3
                st.rerun()

    elif st.session_state.current_step == 3:
        st.header("Step 3: Generating Your Document")

        with st.spinner(
            "Generating comprehensive content based on your outline and feedback..."
        ):
            final_content = doc_gen.generate_content_based_on_outline(
                st.session_state.topic,
                st.session_state.outline,
                st.session_state.user_feedback,
            )

            if final_content:
                st.session_state.final_content = final_content
                st.session_state.current_step = 4
                st.rerun()
    elif st.session_state.current_step == 4:
        st.header("Step 4: Download Your Document")

        col1, col2 = st.columns([1, 1])

        with col1:
            st.subheader("Content Preview")
            with st.expander("View Full Content", expanded=True):
                st.markdown(st.session_state.final_content)

        with col2:
            st.subheader("Download in Multiple Formats")

            with st.spinner("Creating documents in all formats..."):
                multi_gen = MultiFormatDocumentGenerator()
                temp_files = {}

                base_filename = f"{st.session_state.doc_type.lower()}_{st.session_state.topic.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                base_filename = re.sub(r"[^\w\s-]", "", base_filename)

                title = f"{st.session_state.doc_type}: {st.session_state.topic}"

                pdf_path = multi_gen.create_pdf(
                    title,
                    st.session_state.outline,
                    st.session_state.final_content,
                    f"{base_filename}.pdf",
                )
                if pdf_path and os.path.exists(pdf_path):
                    with open(pdf_path, "rb") as f:
                        temp_files["pdf"] = (f.read(), f"{base_filename}.pdf")

                docx_path = multi_gen.create_docx(
                    title,
                    st.session_state.outline,
                    st.session_state.final_content,
                    f"{base_filename}.docx",
                )
                if docx_path and os.path.exists(docx_path):
                    with open(docx_path, "rb") as f:
                        temp_files["docx"] = (f.read(), f"{base_filename}.docx")

                txt_path = multi_gen.create_txt(
                    title,
                    st.session_state.outline,
                    st.session_state.final_content,
                    f"{base_filename}.txt",
                )
                if txt_path and os.path.exists(txt_path):
                    with open(txt_path, "rb") as f:
                        temp_files["txt"] = (f.read(), f"{base_filename}.txt")

                html_path = multi_gen.create_html(
                    title,
                    st.session_state.outline,
                    st.session_state.final_content,
                    f"{base_filename}.html",
                )
                if html_path and os.path.exists(html_path):
                    with open(html_path, "rb") as f:
                        temp_files["html"] = (f.read(), f"{base_filename}.html")

            st.success("All document formats are ready!")

            st.markdown("### Choose Your Format:")

            col1, col2, col3, col4 = st.columns(4)

            with col1:
                if "pdf" in temp_files:
                    st.download_button(
                        label="PDF",
                        data=temp_files["pdf"][0],
                        file_name=temp_files["pdf"][1],
                        mime="application/pdf",
                        use_container_width=True,
                        help="Professional PDF format",
                    )
                else:
                    st.error("PDF generation failed")

            with col2:
                if "docx" in temp_files:
                    st.download_button(
                        label="Word DOCX",
                        data=temp_files["docx"][0],
                        file_name=temp_files["docx"][1],
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        help="Microsoft Word format",
                    )
                else:
                    st.error("DOCX generation failed")

            with col3:
                if "txt" in temp_files:
                    st.download_button(
                        label="Plain Text",
                        data=temp_files["txt"][0],
                        file_name=temp_files["txt"][1],
                        mime="text/plain",
                        use_container_width=True,
                        help="Simple text format",
                    )
                else:
                    st.error("TXT generation failed")

            with col4:
                if "html" in temp_files:
                    st.download_button(
                        label="HTML",
                        data=temp_files["html"][0],
                        file_name=temp_files["html"][1],
                        mime="text/html",
                        use_container_width=True,
                        help="Web page format",
                    )
                else:
                    st.error("HTML generation failed")
            st.markdown("---")
            st.subheader("Format Comparison")

            format_info = {
                "PDF": "Best for printing, professional reports, and formal documents",
                "DOCX": "Editable in Microsoft Word, good for further modifications",
                "TXT": "Universal format, smallest file size, compatible with everything",
                "HTML": "Web-ready, can be viewed in browsers, easy to publish online",
            }

            for format_name, description in format_info.items():
                st.markdown(f"**{format_name}**: {description}")
            for file_path in [pdf_path, docx_path, txt_path, html_path]:
                try:
                    if file_path and os.path.exists(file_path):
                        os.unlink(file_path)
                except:
                    pass

            st.markdown("---")
            st.subheader("Generate Another")
            if st.button("Create New Document", use_container_width=True):
                for key in list(st.session_state.keys()):
                    if key not in ["api_key", "selected_model"]:
                        del st.session_state[key]
                st.session_state.current_step = 1
                st.rerun()

    st.markdown("---")
    st.markdown(
        """
        <div style='text-align: center'>
            <p>Built with ❤️ using Streamlit, LangChain, and Gemini AI | Multi-Format Support</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

if __name__ == "__main__":
    main()
